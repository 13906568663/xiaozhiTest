"""多格式文件解析器，将上传文件统一转为纯文本。

支持格式：txt / markdown / pdf / doc / docx / xlsx / html / csv。
所有解析均在线程池中执行（CPU 密集型操作不阻塞事件循环）。
"""

from __future__ import annotations

import asyncio
import csv
import io
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path, PurePosixPath


@dataclass
class ParseResult:
    content: str
    source_type: str
    metadata: dict = field(default_factory=dict)


_EXTENSION_MAP: dict[str, str] = {
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".pdf": "pdf",
    ".doc": "doc",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".html": "html",
    ".htm": "html",
    ".csv": "csv",
}

SUPPORTED_EXTENSIONS = set(_EXTENSION_MAP.keys())


def _detect_source_type(file_name: str) -> str:
    ext = PurePosixPath(file_name).suffix.lower()
    source_type = _EXTENSION_MAP.get(ext)
    if source_type is None:
        raise ValueError(
            f"不支持的文件格式 '{ext}'，支持：{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return source_type


def _parse_text(data: bytes) -> ParseResult:
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            text = data.decode(encoding)
            return ParseResult(
                content=text, source_type="text", metadata={"encoding": encoding}
            )
        except (UnicodeDecodeError, LookupError):
            continue
    return ParseResult(
        content=data.decode("utf-8", errors="replace"), source_type="text"
    )


def _parse_markdown(data: bytes) -> ParseResult:
    result = _parse_text(data)
    result.source_type = "markdown"
    return result


def _parse_pdf(data: bytes) -> ParseResult:
    import pdfplumber

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
    return ParseResult(
        content="\n\n".join(pages_text),
        source_type="pdf",
        metadata={"page_count": len(pages_text)},
    )


def _parse_docx(data: bytes) -> ParseResult:
    from zipfile import BadZipFile

    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except BadZipFile:
        raise ValueError(
            "文件不是有效的 .docx 格式。"
            "请确认文件未损坏，或尝试用 Word / WPS 重新保存后上传。"
        )
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return ParseResult(
        content="\n\n".join(paragraphs),
        source_type="docx",
        metadata={"paragraph_count": len(paragraphs)},
    )


def _parse_doc(data: bytes) -> ParseResult:
    """Convert legacy .doc (Word 97-2003) to .docx via system tools, then parse.

    Tries LibreOffice (``soffice``) first, then macOS ``textutil``.
    """
    converter = shutil.which("soffice") or shutil.which("textutil")
    if converter is None:
        raise ValueError(
            "解析 .doc 文件需要系统安装 LibreOffice 或在 macOS 上使用 textutil。"
            "建议将文件另存为 .docx 格式后重新上传。"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = Path(tmpdir) / "input.doc"
        doc_path.write_bytes(data)
        docx_path = Path(tmpdir) / "input.docx"

        try:
            if "soffice" in converter:
                subprocess.run(
                    [
                        converter,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmpdir,
                        str(doc_path),
                    ],
                    timeout=60,
                    check=True,
                    capture_output=True,
                )
            else:
                subprocess.run(
                    ["textutil", "-convert", "docx", str(doc_path)],
                    timeout=60,
                    check=True,
                    capture_output=True,
                )
        except (subprocess.SubprocessError, OSError) as exc:
            raise ValueError(
                f"转换 .doc 文件失败：{exc}。"
                "建议将文件另存为 .docx 格式后重新上传。"
            ) from exc

        if not docx_path.exists():
            raise ValueError(
                "转换 .doc 文件后未生成 .docx。"
                "建议将文件另存为 .docx 格式后重新上传。"
            )

        result = _parse_docx(docx_path.read_bytes())
        result.source_type = "doc"
        return result


def _parse_xlsx(data: bytes) -> ParseResult:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets_text: list[str] = []
    sheet_names: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_names.append(sheet_name)
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) for cell in row if cell is not None]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            sheets_text.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))

    wb.close()
    return ParseResult(
        content="\n\n".join(sheets_text),
        source_type="xlsx",
        metadata={"sheet_names": sheet_names, "sheet_count": len(sheet_names)},
    )


def _parse_html(data: bytes) -> ParseResult:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    title = soup.title.string if soup.title else None
    return ParseResult(
        content=text,
        source_type="html",
        metadata={"title": title},
    )


def _parse_csv(data: bytes) -> ParseResult:
    text_content = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text_content))
    rows: list[str] = []
    for row in reader:
        cells = [cell.strip() for cell in row if cell.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return ParseResult(
        content="\n".join(rows),
        source_type="csv",
        metadata={"row_count": len(rows)},
    )


_PARSERS: dict[str, type[None] | None] = {
    "text": None,
    "markdown": None,
    "pdf": None,
    "doc": None,
    "docx": None,
    "xlsx": None,
    "html": None,
    "csv": None,
}

_PARSER_FUNCS = {
    "text": _parse_text,
    "markdown": _parse_markdown,
    "pdf": _parse_pdf,
    "doc": _parse_doc,
    "docx": _parse_docx,
    "xlsx": _parse_xlsx,
    "html": _parse_html,
    "csv": _parse_csv,
}


def _parse_sync(file_bytes: bytes, file_name: str) -> ParseResult:
    source_type = _detect_source_type(file_name)
    parser_func = _PARSER_FUNCS[source_type]
    return parser_func(file_bytes)


async def parse_file(file_bytes: bytes, file_name: str) -> ParseResult:
    """根据文件扩展名自动选择解析器，返回纯文本 + 元数据。

    在线程池中执行以避免阻塞事件循环。
    """
    return await asyncio.to_thread(_parse_sync, file_bytes, file_name)
